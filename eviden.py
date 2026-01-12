import streamlit as st
import pandas as pd
from fpdf import FPDF
import re
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement 
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

# --- KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Sistem LCKB & Evidence", layout="wide", page_icon="🎓")

# --- DATA MASTER ---
BULAN_INDO = {
    1: 'Januari', 2: 'Februari', 3: 'Maret', 4: 'April', 5: 'Mei', 6: 'Juni',
    7: 'Juli', 8: 'Agustus', 9: 'September', 10: 'Oktober', 11: 'November', 12: 'Desember'
}

DAFTAR_DOSEN_RESMI = [
    "Amran Eku, M. Pd.", "Andi Nurmawaddah, M. Pd.", "Asmiraty. M.Pd.I",
    "Dr. Andy, S.Pd.I., M.Pd.", "Dr. Kartini Limatahu, MA", "Dr. Khalid Hasan Minabari, MA.",
    "Dr. M. Ridha Assagaf, S. Ag., M.Pd.", "Dr. Mubin Noho, S.Ag., M.Ag", "Dr. Usman Ilyas, M. Pd.",
    "Dra. Nursin Sapil, M.Pd.I", "Drs. Hi Ibrahim Muhammad, M.Pd.I", "Drs. Kamarun M. Sebe, M.Pd.",
    "Drs. Ramli Yusuf, M.Pd.", "Elfira Mahmud, M.Pd.", "Hamdy M. Zen, M. Pd.I",
    "Irno, S.Hum., M.Hum.", "M. Rizki Hi. Aman", "Mawardi Djamaludin",
    "Minggusta Juliadarma, M.Pd.I", "Mudayanah, M.Pd", "Puji Dwi Rahayu, M.Pd.",
    "Rinelsa R. Husaen, M.Pd", "Yani Djawa, S.Pd,M.Pd.Si", "Zainuddin Arifin"
]

KATEGORI_LABEL = {
    'A': 'BIDANG PENDIDIKAN DAN PENGAJARAN',
    'B': 'BIDANG PENELITIAN',
    'C': 'BIDANG PENGABDIAN MASYARAKAT',
    'D': 'BIDANG PENUNJANG AKADEMIK'
}

# --- GOOGLE DRIVE MANAGER (UPDATED FOR SHARED DRIVE SUPPORT) ---
def get_drive_service():
    if "gcp_service_account" not in st.secrets:
        return None, "Secret 'gcp_service_account' tidak ditemukan di .streamlit/secrets.toml"
    try:
        creds = service_account.Credentials.from_service_account_info(
            st.secrets["gcp_service_account"], scopes=['https://www.googleapis.com/auth/drive']
        )
        return build('drive', 'v3', credentials=creds), None
    except Exception as e:
        return None, str(e)

def get_or_create_folder(service, folder_name, parent_id):
    try:
        # Tambahkan supportsAllDrives=True dan includeItemsFromAllDrives=True
        query = f"mimeType='application/vnd.google-apps.folder' and name='{folder_name}' and '{parent_id}' in parents and trashed=false"
        results = service.files().list(
            q=query, 
            fields="files(id, webViewLink)", 
            supportsAllDrives=True, 
            includeItemsFromAllDrives=True
        ).execute()
        
        files = results.get('files', [])
        if files: 
            return files[0]['id'], files[0]['webViewLink']
        else:
            metadata = {
                'name': folder_name, 
                'mimeType': 'application/vnd.google-apps.folder', 
                'parents': [parent_id]
            }
            # Tambahkan supportsAllDrives=True saat create
            folder = service.files().create(
                body=metadata, 
                fields="id, webViewLink", 
                supportsAllDrives=True
            ).execute()
            return folder['id'], folder['webViewLink']
    except Exception as e:
        raise Exception(f"Gagal buat folder '{folder_name}': {str(e)}")

def upload_file_to_drive(file_obj, filename, dosen_name, tahun, semester, kategori, nama_kegiatan):
    service, err = get_drive_service()
    if not service: return None, err
    
    try:
        root_id = st.secrets["target_folder_id"]
        
        # 1. Buat Struktur Folder (Support Shared Drive)
        dosen_id, _ = get_or_create_folder(service, dosen_name, root_id)
        tahun_id, _ = get_or_create_folder(service, str(tahun), dosen_id)
        sem_id, _ = get_or_create_folder(service, semester, tahun_id)
        kat_clean = KATEGORI_LABEL[kategori].replace("BIDANG ", "")
        kat_id, _ = get_or_create_folder(service, kat_clean, sem_id)
        
        # 2. Buat Folder Kegiatan
        safe_kegiatan = re.sub(r'[\\/*?:"<>|]', "", nama_kegiatan)[:50] 
        kegiatan_id, kegiatan_link = get_or_create_folder(service, safe_kegiatan, kat_id)
        
        # 3. Upload File (Support Shared Drive)
        media = MediaIoBaseUpload(file_obj, mimetype=file_obj.type)
        file_meta = {'name': filename, 'parents': [kegiatan_id]}
        
        service.files().create(
            body=file_meta, 
            media_body=media, 
            fields='webViewLink', 
            supportsAllDrives=True
        ).execute()
        
        return kegiatan_link, None
    except Exception as e:
        return None, str(e)

# --- FUNGSI HELPER & PARSING ---
def normalize_name(raw_name):
    if pd.isna(raw_name): return ""
    name = str(raw_name).upper()
    gelar_pattern = r'\b(DR|DRA|DRS|IR|S\. ?PD|M\. ?PD|S\. ?AG|M\. ?AG|S\. ?HUM|M\. ?HUM|S\. ?SI|M\. ?SI|S\. ?KOM|M\. ?KOM|PH\. ?D|M\. ?PI|S\. ?H|M\. ?H|I|II|S\. ?SOS|M\. ?SOS|M\. ?A|M\. ?PHIL|M\. ?PD\. ?I|S\. ?PD\. ?I)\b'
    name = re.sub(gelar_pattern, '', name)
    name = re.sub(r'[.,]', ' ', name)
    name = " ".join(name.split())
    return name

def extract_drive_id(url):
    if not isinstance(url, str): return None
    patterns = [r'/d/([a-zA-Z0-9_-]{25,})', r'id=([a-zA-Z0-9_-]{25,})', r'open\?id=([a-zA-Z0-9_-]{25,})']
    for p in patterns:
        m = re.search(p, url)
        if m: return m.group(1)
    return None

def process_links(raw_link_str):
    if pd.isna(raw_link_str) or not isinstance(raw_link_str, str): return []
    raw_links = re.split(r'[,\n\s]+', raw_link_str)
    processed = []
    for link in raw_links:
        link = link.strip().replace('"', '').replace("'", "")
        if len(link) < 10: continue
        fid = extract_drive_id(link)
        thumb = f"https://lh3.googleusercontent.com/d/{fid}=s400" if fid else None
        processed.append({'original': link, 'thumb': thumb})
    return processed

def parse_evidence_full(row):
    jenis = str(row.get('Pilih Jenis Ujian', ''))
    
    def find_val(keywords):
        for c in row.index:
            if all(k.lower() in c.lower() for k in keywords):
                return row[c]
        return None

    raw_ba = raw_foto = raw_naskah = raw_undangan = raw_penunjukan = None

    if 'UAS' in jenis:
        raw_ba = find_val(['berita', 'acara', 'uas'])
        raw_foto = find_val(['foto', 'dokumentasi', 'uas'])
        raw_naskah = find_val(['naskah', 'soal'])
    elif 'Kompre' in jenis:
        raw_ba = find_val(['berita', 'acara', 'kompre'])
        raw_foto = find_val(['foto', 'dokumentasi', 'kompre'])
        raw_penunjukan = find_val(['penunjukan', 'penguji']) 
    else: # Proposal/Skripsi
        raw_ba = find_val(['berita', 'acara'])
        raw_foto = find_val(['foto', 'dokumentasi'])
        raw_undangan = find_val(['undangan']) 

    return {
        'ba': process_links(raw_ba), 
        'foto': process_links(raw_foto), 
        'naskah': process_links(raw_naskah),
        'undangan': process_links(raw_undangan),
        'penunjukan': process_links(raw_penunjukan)
    }

@st.cache_data
def load_data(url):
    try:
        df = pd.read_csv(url)
        df.columns = df.columns.str.strip() 
        df['Timestamp'] = pd.to_datetime(df['Timestamp'], dayfirst=True, errors='coerce')
        df['Bulan'] = df['Timestamp'].dt.month
        df['Tahun'] = df['Timestamp'].dt.year
        keywords = ['dosen', 'pembimbing', 'penguji']
        target_cols = [c for c in df.columns if 'nama' in c.lower() and any(k in c.lower() for k in keywords)]
        return df, target_cols
    except Exception as e:
        st.error(f"Error membaca CSV: {e}")
        return None, None

# --- GENERATOR PDF EVIDENCE ---
class EvidencePDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 5, 'KEMENTERIAN AGAMA REPUBLIK INDONESIA', 0, 1, 'C')
        self.cell(0, 5, 'INSTITUT AGAMA ISLAM NEGERI (IAIN) TERNATE', 0, 1, 'C')
        self.ln(2); self.line(10, 22, 200, 22); self.ln(5)

def create_evidence_pdf_bytes(df_filtered, dosen_name, periode_label):
    try:
        pdf = EvidencePDF('P', 'mm', 'A4'); pdf.add_page()
        pdf.set_font('Arial', 'B', 11); pdf.cell(0, 6, 'LAPORAN BUKTI FISIK PELAKSANAAN UJIAN', 0, 1, 'C')
        pdf.set_font('Arial', '', 10); pdf.cell(0, 5, f'PERIODE: {periode_label.upper()}', 0, 1, 'C'); pdf.ln(8)
        pdf.set_font('Arial', '', 10); pdf.cell(30, 5, 'Nama Dosen', 0, 0); pdf.cell(5, 5, ':', 0, 0); pdf.cell(0, 5, dosen_name, 0, 1); pdf.ln(5)

        # BAGIAN 1: UAS
        df_uas = df_filtered[df_filtered['Pilih Jenis Ujian'].str.contains('UAS', case=False, na=False)]
        if not df_uas.empty:
            pdf.set_font('Arial', 'B', 10); pdf.cell(0, 8, 'A. UJIAN AKHIR SEMESTER (UAS)', 0, 1, 'L')
            pdf.set_fill_color(230, 230, 230); pdf.set_font('Arial', 'B', 8)
            cols = [('NO', 10), ('MATA KULIAH / KELAS', 75), ('BERITA ACARA', 35), ('DOK', 35), ('SOAL', 35)]
            for txt, w in cols: pdf.cell(w, 8, txt, 1, 0, 'C', 1)
            pdf.ln(); pdf.set_font('Arial', '', 8); no = 1
            for _, row in df_uas.iterrows():
                ev = parse_evidence_full(row)
                matkul = f"{row.get('Nama Matkul','-')} ({row.get('Nama Kelas','-')})".encode('latin-1', 'ignore').decode('latin-1')
                l_ba = ev['ba'][0]['original'] if ev['ba'] else ""
                l_dok = ev['foto'][0]['original'] if ev['foto'] else ""
                l_soal = ev['naskah'][0]['original'] if ev['naskah'] else ""
                
                if pdf.get_y() + 8 > 260: pdf.add_page()
                pdf.cell(10, 8, str(no), 1, 0, 'C')
                pdf.cell(75, 8, matkul[:40], 1, 0, 'L')
                pdf.cell(35, 8, "Buka File" if l_ba else "-", 1, 0, 'C', link=l_ba)
                pdf.cell(35, 8, "Buka File" if l_dok else "-", 1, 0, 'C', link=l_dok)
                pdf.cell(35, 8, "Buka File" if l_soal else "-", 1, 1, 'C', link=l_soal)
                no += 1
            pdf.ln(5)

        # BAGIAN 2: NON-UAS
        df_non = df_filtered[~df_filtered['Pilih Jenis Ujian'].str.contains('UAS', case=False, na=False)]
        if not df_non.empty:
            pdf.set_font('Arial', 'B', 10); pdf.cell(0, 8, 'B. UJIAN LAINNYA', 0, 1, 'L')
            pdf.set_fill_color(230, 230, 230); pdf.set_font('Arial', 'B', 8)
            cols = [('NO', 10), ('URAIAN KEGIATAN', 75), ('BERITA ACARA', 35), ('SURAT/SK', 35), ('DOK', 35)]
            for txt, w in cols: pdf.cell(w, 8, txt, 1, 0, 'C', 1)
            pdf.ln(); pdf.set_font('Arial', '', 8); no = 1
            for _, row in df_non.iterrows():
                ev = parse_evidence_full(row)
                ur = f"{row.get('Pilih Jenis Ujian')} - {row.get('Nama Lengkap Mahasiswa','-')}".encode('latin-1', 'ignore').decode('latin-1')
                l_ba = ev['ba'][0]['original'] if ev['ba'] else ""
                l_surat = ""
                if ev['undangan']: l_surat = ev['undangan'][0]['original']
                elif ev['penunjukan']: l_surat = ev['penunjukan'][0]['original']
                l_dok = ev['foto'][0]['original'] if ev['foto'] else ""
                
                if pdf.get_y() + 8 > 260: pdf.add_page()
                pdf.cell(10, 8, str(no), 1, 0, 'C'); pdf.cell(75, 8, ur[:45], 1, 0, 'L')
                pdf.cell(35, 8, "Buka File" if l_ba else "-", 1, 0, 'C', link=l_ba)
                pdf.cell(35, 8, "Buka File" if l_surat else "-", 1, 0, 'C', link=l_surat)
                pdf.cell(35, 8, "Buka File" if l_dok else "-", 1, 1, 'C', link=l_dok)
                no += 1

        pdf.ln(10)
        if pdf.get_y() > 240: pdf.add_page()
        pdf.set_x(120); pdf.cell(60, 5, f'Ternate, {datetime.now().strftime("%d-%m-%Y")}', 0, 1, 'C')
        pdf.set_x(120); pdf.cell(60, 5, 'Dosen Yang Melaporkan,', 0, 1, 'C'); pdf.ln(20)
        pdf.set_x(120); pdf.set_font('Arial', 'B', 9); pdf.cell(60, 5, dosen_name, 0, 1, 'C')
        return pdf.output(dest='S').encode('latin-1', 'ignore')
    except: return None

# --- GENERATOR WORD EVIDENCE ---
def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    part = paragraph.part; r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    if color: c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
    if underline: u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr); new_text = OxmlElement('w:t'); new_text.text = text; new_run.append(new_text)
    hyperlink.append(new_run); paragraph._element.append(hyperlink); return hyperlink

def create_evidence_docx_bytes(df_filtered, dosen_name, periode_label):
    try:
        doc = Document()
        doc.add_paragraph('LAPORAN BUKTI FISIK PELAKSANAAN UJIAN').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'PERIODE: {periode_label}').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'Nama Dosen: {dosen_name}\n')

        def fill_cell(cell, link, text_ok="Buka File"):
            if link: add_hyperlink(cell.add_paragraph(), link, text_ok)
            else: cell.text = "-"

        df_uas = df_filtered[df_filtered['Pilih Jenis Ujian'].str.contains('UAS', case=False, na=False)]
        if not df_uas.empty:
            doc.add_paragraph('A. UJIAN AKHIR SEMESTER (UAS)').runs[0].bold = True
            table = doc.add_table(rows=1, cols=5); table.style = 'Table Grid'
            hdr = table.rows[0].cells; 
            hdr[0].text='NO'; hdr[1].text='MATA KULIAH'; hdr[2].text='BA'; hdr[3].text='DOK'; hdr[4].text='SOAL'
            no = 1
            for _, row in df_uas.iterrows():
                ev = parse_evidence_full(row)
                rc = table.add_row().cells; rc[0].text=str(no); rc[1].text=f"{row.get('Nama Matkul','-')} ({row.get('Nama Kelas','-')})"
                fill_cell(rc[2], ev['ba'][0]['original'] if ev['ba'] else None)
                fill_cell(rc[3], ev['foto'][0]['original'] if ev['foto'] else None)
                fill_cell(rc[4], ev['naskah'][0]['original'] if ev['naskah'] else None)
                no += 1
            doc.add_paragraph('\n')

        df_non = df_filtered[~df_filtered['Pilih Jenis Ujian'].str.contains('UAS', case=False, na=False)]
        if not df_non.empty:
            doc.add_paragraph('B. UJIAN LAINNYA').runs[0].bold = True
            table = doc.add_table(rows=1, cols=5); table.style = 'Table Grid'
            hdr = table.rows[0].cells; hdr[0].text='NO'; hdr[1].text='URAIAN'; hdr[2].text='BA'; hdr[3].text='SURAT/SK'; hdr[4].text='DOK'
            no = 1
            for _, row in df_non.iterrows():
                ev = parse_evidence_full(row)
                rc = table.add_row().cells; rc[0].text=str(no); rc[1].text=f"{row.get('Pilih Jenis Ujian')} - {row.get('Nama Lengkap Mahasiswa')}"
                l_surat = None
                if ev['undangan']: l_surat = ev['undangan'][0]['original']
                elif ev['penunjukan']: l_surat = ev['penunjukan'][0]['original']
                fill_cell(rc[2], ev['ba'][0]['original'] if ev['ba'] else None)
                fill_cell(rc[3], l_surat)
                fill_cell(rc[4], ev['foto'][0]['original'] if ev['foto'] else None)
                no += 1
        
        doc.add_paragraph('\n')
        sig = doc.add_paragraph(f'Ternate, {datetime.now().strftime("%d-%m-%Y")}\nDosen Yang Melaporkan,\n\n\n\n\n{dosen_name}')
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f = BytesIO(); doc.save(f); return f.getvalue()
    except: return None

# --- GENERATOR LCKB (LCKB FIXES) ---
class LCKB_PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 5, 'KEMENTERIAN AGAMA REPUBLIK INDONESIA', 0, 1, 'C')
        self.cell(0, 5, 'INSTITUT AGAMA ISLAM NEGERI (IAIN) TERNATE', 0, 1, 'C')
        self.ln(2); self.line(10, 22, 200, 22); self.ln(5)

def create_lckb_pdf_bytes(data_items, dosen_name, bulan, tahun, nama_dekan, nip_dekan):
    try:
        pdf = LCKB_PDF('P', 'mm', 'A4'); pdf.add_page()
        pdf.set_font('Arial', 'B', 11); pdf.cell(0, 6, 'LAPORAN CAPAIAN KINERJA BULANAN (LCKB)', 0, 1, 'C')
        pdf.set_font('Arial', '', 10); pdf.cell(0, 5, f'BULAN: {str(bulan).upper()} {tahun}', 0, 1, 'C'); pdf.ln(5)
        pdf.cell(30, 5, 'Nama', 0, 0); pdf.cell(5, 5, ':', 0, 0); pdf.cell(0, 5, dosen_name, 0, 1); pdf.ln(5)
        pdf.set_fill_color(230, 230, 230); pdf.set_font('Arial', 'B', 9)
        cols = [('NO', 10), ('URAIAN TUGAS/KEGIATAN', 80), ('VOL', 15), ('SAT', 25), ('BUKTI FISIK (KLIK FILE)', 60)]
        for txt, w in cols: pdf.cell(w, 8, txt, 1, 0, 'C', 1)
        pdf.ln(8); pdf.set_font('Arial', '', 8)
        
        for kat in ['A', 'B', 'C', 'D']:
            pdf.set_font('Arial', 'B', 9); pdf.cell(190, 7, f"{kat}. {KATEGORI_LABEL[kat]}", 1, 1, 'L', 1)
            pdf.set_font('Arial', '', 8)
            items = [x for x in data_items if x['kategori'] == kat]
            
            if not items: 
                pdf.cell(190, 6, " (Tidak ada kegiatan)", 1, 1, 'C')
            else:
                no = 1 
                for item in items:
                    desc = item['uraian'].encode('latin-1', 'ignore').decode('latin-1')
                    pdf.cell(10, 6, str(no), 1, 0, 'C'); pdf.cell(80, 6, desc[:50], 1, 0, 'L')
                    pdf.cell(15, 6, str(item['volume']), 1, 0, 'C'); pdf.cell(25, 6, item['satuan'], 1, 0, 'C')
                    
                    # LOGIC PRINT LINK
                    links_raw = item.get('bukti_list', [])
                    if links_raw and isinstance(links_raw, list):
                        w_cell = 60
                        count = len(links_raw)
                        w_item = w_cell / count if count > 0 else 60
                        pdf.set_font('Arial', 'U', 8); pdf.set_text_color(0, 0, 255)
                        for lnk in links_raw:
                            lbl = lnk['label']; url = lnk['url']
                            if url: pdf.cell(w_item, 6, lbl, 1, 0, 'C', link=url)
                            else:
                                pdf.set_text_color(0,0,0); pdf.set_font('Arial', '', 8)
                                pdf.cell(w_item, 6, "-", 1, 0, 'C')
                                pdf.set_font('Arial', 'U', 8); pdf.set_text_color(0, 0, 255)
                        pdf.set_text_color(0, 0, 0); pdf.set_font('Arial', '', 8)
                        pdf.ln()
                    else:
                        link_val = str(item['bukti']).strip()
                        txt_display = "Link Folder" if "http" in link_val else link_val
                        if "http" in link_val: pdf.cell(60, 6, txt_display, 1, 1, 'C', link=link_val)
                        else: pdf.cell(60, 6, txt_display, 1, 1, 'C')
                    no += 1
        
        pdf.ln(10); y = pdf.get_y(); pdf.set_xy(20, y); pdf.cell(60, 5, "Mengetahui, Dekan FTIK,", 0, 0, 'C')
        pdf.set_xy(120, y); pdf.cell(60, 5, f'Ternate, {datetime.now().strftime("%d-%m-%Y")}', 0, 1, 'C')
        pdf.set_x(120); pdf.cell(60, 5, 'Yang Melaporkan,', 0, 1, 'C'); pdf.ln(15)
        pdf.set_font('Arial', 'B', 9); pdf.set_xy(20, pdf.get_y()); pdf.cell(60, 5, nama_dekan, 0, 0, 'C')
        pdf.set_xy(120, pdf.get_y()); pdf.cell(60, 5, dosen_name, 0, 1, 'C')
        return pdf.output(dest='S').encode('latin-1', 'ignore')
    except: return None

def create_lckb_docx_bytes(data_items, dosen_name, bulan, tahun, nama_dekan, nip_dekan):
    try:
        doc = Document(); doc.add_paragraph('LAPORAN LCKB').alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = doc.add_table(rows=1, cols=5); table.style = 'Table Grid'
        for i, txt in enumerate(['NO', 'URAIAN', 'VOL', 'SAT', 'BUKTI FISIK']): table.rows[0].cells[i].text = txt
        
        for kat in ['A', 'B', 'C', 'D']:
            row = table.add_row().cells; row[0].merge(row[4]); row[0].text = f"{kat}. {KATEGORI_LABEL[kat]}"
            no = 1
            for it in [x for x in data_items if x['kategori'] == kat]:
                cells = table.add_row().cells; cells[0].text, cells[1].text, cells[2].text, cells[3].text = str(no), it['uraian'], str(it['volume']), it['satuan']
                
                links_raw = it.get('bukti_list', [])
                cell_bukti = cells[4]
                if links_raw and isinstance(links_raw, list):
                    p = cell_bukti.add_paragraph()
                    valid_links = [l for l in links_raw if l['url']]
                    if valid_links:
                        for i, lnk in enumerate(valid_links):
                            add_hyperlink(p, lnk['url'], lnk['label'])
                            if i < len(valid_links) - 1: p.add_run(" | ")
                    else: cell_bukti.text = "-"
                else:
                    link_val = str(it['bukti']).strip()
                    if "http" in link_val: add_hyperlink(cell_bukti.add_paragraph(), link_val, "Link Folder")
                    else: cell_bukti.text = link_val
                no += 1
        f = BytesIO(); doc.save(f); return f.getvalue()
    except: return None

# --- MAIN APP ---
url = "https://docs.google.com/spreadsheets/d/e/2PACX-1vQinSdwQBQZj649QKRimqqmTFQ0WaSlEHucehHOEg7jvTaioDXe0snCcpo3kTJJsnFrIcqEasjif9E8/pub?output=csv"
df, target_cols = load_data(url)

if 'manual_data' not in st.session_state: st.session_state['manual_data'] = []
if 'last_dosen' not in st.session_state: st.session_state['last_dosen'] = ""

st.sidebar.title("Navigasi")
menu = st.sidebar.radio("Menu:", ["1. Cek Evidence & Cetak", "2. Buat LCKB (Dosen)"])
nama_dekan = st.sidebar.text_input("Nama Dekan", "Dr. H. Sahjad M. Aksan, M.Phil")
nip_dekan = st.sidebar.text_input("NIP Dekan", "19xxxxxxx")

if df is not None:
    # --- MENU 1 ---
    if menu == "1. Cek Evidence & Cetak":
        st.title("📂 Data Evidence")
        c1, c2, c3 = st.columns(3)
        dsn = c1.selectbox("Dosen:", DAFTAR_DOSEN_RESMI)
        mode = c2.selectbox("Filter Waktu:", ["Bulanan", "Semester Ganjil", "Semester Genap", "Tahunan", "Semua Data"])
        thn = c3.number_input("Tahun", 2024, 2030, datetime.now().year)
        
        df_d = df[df.astype(str).apply(lambda x: x.str.contains(normalize_name(dsn), case=False)).any(axis=1)].copy()
        label = f"TAHUN {thn}"
        if mode == "Bulanan":
            bln = st.selectbox("Bulan:", list(BULAN_INDO.values()))
            b_int = list(BULAN_INDO.keys())[list(BULAN_INDO.values()).index(bln)]
            df_f = df_d[(df_d['Bulan']==b_int)&(df_d['Tahun']==thn)]; label = f"{bln.upper()} {thn}"
        elif mode == "Semester Ganjil":
            df_f = df_d[(df_d['Bulan'].isin([7,8,9,10,11,12]))&(df_d['Tahun']==thn)]; label = f"SEMESTER GANJIL {thn}"
        elif mode == "Semester Genap":
            df_f = df_d[(df_d['Bulan'].isin([1,2,3,4,5,6]))&(df_d['Tahun']==thn)]; label = f"SEMESTER GENAP {thn}"
        elif mode == "Tahunan":
            df_f = df_d[df_d['Tahun']==thn]
        else: df_f = df_d; label = "SEMUA RIWAYAT DATA"

        st.divider(); st.write(f"Menampilkan **{len(df_f)}** data ({label})")
        tab1, tab2 = st.tabs(["Preview", "Download"])
        with tab1:
            if df_f.empty: st.warning("Data Kosong")
            for _, r in df_f.iterrows():
                ev = parse_evidence_full(r)
                with st.expander(f"{r['Timestamp'].strftime('%d %b')} | {r.get('Pilih Jenis Ujian')} | {r.get('Nama Lengkap Mahasiswa','-')}"):
                    ca, cb = st.columns([1,2]); 
                    if ev['foto']: ca.image([x['thumb'] for x in ev['foto'] if x['thumb']], width=100)
                    for k in ['ba','undangan','penunjukan','naskah','foto']: 
                        if ev[k]: cb.write(f"**{k.upper()}:** " + ", ".join([f"[{x['original']}]({x['original']})" for x in ev[k]]))
        with tab2:
            if not df_f.empty:
                c_p, c_w = st.columns(2)
                c_p.download_button("📄 PDF Laporan", create_evidence_pdf_bytes(df_f, dsn, label), f"Lap_{dsn}.pdf", "application/pdf")
                c_w.download_button("📝 Word Laporan", create_evidence_docx_bytes(df_f, dsn, label), f"Lap_{dsn}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # --- MENU 2 ---
    elif menu == "2. Buat LCKB (Dosen)":
        st.title("📝 Buat LCKB (Upload ke Drive)")
        c1, c2, c3 = st.columns(3)
        dsn = c1.selectbox("Dosen:", DAFTAR_DOSEN_RESMI)
        
        # LOGIKA RESET JIKA GANTI DOSEN (RESET FIX)
        if st.session_state['last_dosen'] != dsn:
            st.session_state['manual_data'] = [] 
            st.session_state['last_dosen'] = dsn

        bln = c2.selectbox("Bulan:", list(BULAN_INDO.values()))
        thn = c3.number_input("Tahun", 2024, 2030, datetime.now().year)
        b_int = list(BULAN_INDO.keys())[list(BULAN_INDO.values()).index(bln)]
        semester = "Semester Ganjil" if b_int >= 7 else "Semester Genap"
        
        mask_d = pd.Series(False, index=df.index)
        for c in target_cols: mask_d |= df[c].apply(normalize_name).str.contains(normalize_name(dsn), na=False)
        df_auto = df[mask_d & (df['Bulan']==b_int) & (df['Tahun']==thn)]
        
        with st.expander("➕ Tambah Kegiatan & Upload Bukti", expanded=True):
            with st.form("upload_form"):
                kat = st.selectbox("Kategori", list(KATEGORI_LABEL.keys()), format_func=lambda x:KATEGORI_LABEL[x])
                ur = st.text_input("Uraian Kegiatan (Nama Folder)", placeholder="Contoh: Mengajar MK Fiqih Kelas A")
                col_vol, col_sat = st.columns(2)
                vol = col_vol.number_input("Volume", 1); sat = col_sat.text_input("Satuan", "SKS/Kegiatan")
                st.markdown("---")
                uploaded_file = st.file_uploader("Upload Bukti (PDF/Gambar)")
                
                # UPLOAD + ERROR HANDLING (NO RERUN ON ERROR)
                if st.form_submit_button("Simpan & Upload"):
                    if not ur: st.error("Uraian Kegiatan wajib diisi.")
                    else:
                        folder_link, err_msg = None, None
                        if uploaded_file:
                            with st.spinner("Sedang proses upload..."):
                                folder_link, err_msg = upload_file_to_drive(uploaded_file, uploaded_file.name, dsn, thn, semester, kat, ur)
                        
                        if err_msg: st.error(f"❌ Upload Gagal: {err_msg}")
                        else:
                            if uploaded_file: st.success("✅ Upload Berhasil!")
                            st.session_state['manual_data'].append({
                                'kategori': kat, 'uraian': ur, 'volume': vol, 'satuan': sat, 'bukti': folder_link if folder_link else "-"
                            })
                            st.rerun()

        st.divider(); st.subheader("📋 Draft Laporan")
        
        # 1. AUTO DATA (SATUAN FIX KELAS/MHS)
        if not df_auto.empty:
            st.info("Data Otomatis (Ujian/Evidence):")
            auto_data = []
            for _, r in df_auto.iterrows():
                ev = parse_evidence_full(r)
                link_collection = []
                if ev['ba']: link_collection.append({'label': 'BA', 'url': ev['ba'][0]['original']})
                if ev['naskah']: link_collection.append({'label': 'Soal', 'url': ev['naskah'][0]['original']})
                if ev['undangan']: link_collection.append({'label': 'Undangan', 'url': ev['undangan'][0]['original']})
                if ev['penunjukan']: link_collection.append({'label': 'SK', 'url': ev['penunjukan'][0]['original']})
                if ev['foto']: link_collection.append({'label': 'Foto', 'url': ev['foto'][0]['original']})
                if not link_collection: link_collection.append({'label': '-', 'url': ''})

                uraian_txt = ""; satuan_txt = "Mhs"
                jenis = str(r.get('Pilih Jenis Ujian', ''))
                
                if 'UAS' in jenis: 
                    uraian_txt = f"Menguji UAS - {r.get('Nama Matkul','-')} ({r.get('Nama Kelas','-')})"
                    satuan_txt = "Kelas" # SATUAN FIX
                else: 
                    uraian_txt = f"Menguji {jenis} - {r.get('Nama Lengkap Mahasiswa','-')}"

                auto_data.append({
                    'kategori': 'A', 'uraian': uraian_txt, 'volume': 1, 'satuan': satuan_txt, 
                    'bukti': 'Auto', 'bukti_list': link_collection
                })
            st.dataframe(pd.DataFrame(auto_data)[['uraian','volume','satuan']], use_container_width=True)
        else:
            auto_data = []

        # 2. MANUAL DATA
        if st.session_state['manual_data']:
            st.warning("Data Manual:")
            df_manual = pd.DataFrame(st.session_state['manual_data'])
            edited_df = st.data_editor(df_manual, num_rows="dynamic", use_container_width=True, key="editor_manual")
            if not df_manual.equals(edited_df):
                st.session_state['manual_data'] = edited_df.to_dict('records')
                st.rerun()
        
        final_list = auto_data + st.session_state['manual_data']
        if final_list:
            st.divider()
            ca, cb = st.columns(2)
            ca.download_button("📄 PDF LCKB", create_lckb_pdf_bytes(final_list, dsn, bln, thn, nama_dekan, nip_dekan), f"LCKB_{bln}.pdf", "application/pdf")
            cb.download_button("📝 Word LCKB", create_lckb_docx_bytes(final_list, dsn, bln, thn, nama_dekan, nip_dekan), f"LCKB_{bln}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")